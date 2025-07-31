import os
import sys
import time
from typing import Optional, List, Dict
import requests
from base64 import b64encode
from html import unescape
import json
from datetime import datetime
from azure.ai.projects import AIProjectClient
from azure.identity import DefaultAzureCredential
from azure.ai.agents import AgentsClient
from azure.ai.agents.models import CodeInterpreterTool, BingGroundingTool, DeepResearchTool, MessageRole
from azure.storage.blob import BlobServiceClient
from docx import Document
from docx.shared import Inches
from io import BytesIO
import markdown
from dotenv import load_dotenv
from pathlib import Path

# Load .env from project root
env_path = Path(__file__).parent.parent / '.env'
load_dotenv(dotenv_path=env_path)

class AzureDevOpsMCPTool:
    def __init__(self, organization: str = None, project: str = None):
        self.organization = organization or "gpsuscodewith"
        self.project = project or "Code-With Engagement Portfolio"
        self.api_version = "7.1-preview.3"
        self.base_url = f"https://dev.azure.com/{self.organization}/{self.project}/_apis"

    def _get_auth_header(self):
        pat = os.getenv("ADO_PAT")
        if not pat:
            raise ValueError("ADO_PAT environment variable is not set")
        return {"Authorization": f"Basic {b64encode(f':{pat}'.encode()).decode()}"}

    def get_azure_devops_story(self, story_id: str) -> str:
        """Get a user story's title, description, state, dates, and assignee from Azure DevOps.
        
        Args:
            story_id: The ID of the Azure DevOps user story to retrieve.
            
        Returns:
            str: Formatted summary of the user story details.
        """
        url = f"{self.base_url}/wit/workitems/{story_id}?api-version={self.api_version}"
        try:
            resp = requests.get(url, headers=self._get_auth_header())
            resp.raise_for_status()
            fields = resp.json()["fields"]

            title = fields.get("System.Title", "No title")
            state = fields.get("System.State", "Unknown")
            assigned_to = fields.get("System.AssignedTo", {}).get("displayName", "Unassigned")
            created = fields.get("System.CreatedDate", "N/A")
            changed = fields.get("System.ChangedDate", "N/A")
            start_date = fields.get("Microsoft.VSTS.Scheduling.StartDate", "N/A")
            finish_date = fields.get("Microsoft.VSTS.Scheduling.FinishDate", "N/A")
            priority = fields.get("Microsoft.VSTS.Common.Priority", "N/A")
            acceptance = fields.get("Microsoft.VSTS.Common.AcceptanceCriteria", "N/A")
            raw_description = fields.get("System.Description", "No description")

            try:
                from bs4 import BeautifulSoup
                description = unescape(BeautifulSoup(raw_description, "html.parser").get_text(strip=True))
            except ImportError:
                import re
                description = unescape(re.sub('<[^<]+?>', '', raw_description))

            summary = (
                f"**Story {story_id}: {title}**\n"
                f"- **State:** {state}\n"
                f"- **Assigned To:** {assigned_to}\n"
                f"- **Created:** {created}\n"
                f"- **Last Modified:** {changed}\n"
                f"- **Start Date:** {start_date}\n"
                f"- **Finish Date:** {finish_date}\n"
                f"- **Priority:** {priority}\n"
                f"- **Acceptance Criteria:** {acceptance}\n"
                f"- **Description:**\n{description}"
            )
            return summary

        except Exception as e:
            return f"Failed to retrieve story {story_id}: {e}"


class DeepResearchMCPTool:
    def __init__(self):
        self.research_templates = {
            "technical": """
            Perform a comprehensive technical analysis for {project_name} project:
            
            1. **Architecture Patterns**: Research best practices for {technology_stack} architecture patterns, scalability considerations, and design principles
            2. **Technology Stack Analysis**: Analyze current trends, performance characteristics, and ecosystem maturity for the proposed technologies
            3. **Implementation Approaches**: Compare different implementation strategies, their trade-offs, and recommended patterns
            4. **Security Considerations**: Identify security best practices, common vulnerabilities, and compliance requirements
            5. **Performance & Scalability**: Research performance optimization techniques and scalability patterns
            6. **Integration Patterns**: Analyze integration approaches with existing systems and third-party services
            
            Focus on actionable insights and specific recommendations for the project context: {project_context}
            """,
            
            "market": """
            Conduct comprehensive market research for {project_name}:
            
            1. **Market Landscape**: Analyze the current market size, growth trends, and key segments
            2. **Competitive Analysis**: Identify direct and indirect competitors, their strengths, weaknesses, and market positioning
            3. **Industry Trends**: Research emerging trends, technologies, and market shifts affecting the domain
            4. **Customer Insights**: Analyze target customer segments, pain points, and preferences
            5. **Regulatory Environment**: Identify relevant regulations, compliance requirements, and industry standards
            6. **Market Opportunities**: Highlight gaps, opportunities, and potential differentiation strategies
            
            Project context: {project_context}
            Target domain: {domain}
            """,
            
            "risk": """
            Perform comprehensive risk assessment for {project_name}:
            
            1. **Technical Risks**: Identify technology-related risks, dependencies, and potential failure points
            2. **Security Risks**: Analyze security threats, vulnerabilities, and compliance risks
            3. **Operational Risks**: Assess deployment, maintenance, and operational challenges
            4. **Business Risks**: Identify market, financial, and strategic risks
            5. **Regulatory Risks**: Analyze compliance, legal, and regulatory challenges
            6. **Mitigation Strategies**: Provide specific risk mitigation approaches and contingency plans
            
            Project details: {project_context}
            Industry: {industry}
            Timeline: {timeline}
            """,
            
            "stakeholder": """
            Conduct stakeholder analysis for {project_name}:
            
            1. **User Personas**: Research and define primary user segments, their needs, behaviors, and pain points
            2. **Business Stakeholders**: Identify key business stakeholders, their priorities, and success criteria
            3. **Technical Stakeholders**: Analyze technical team requirements, constraints, and preferences
            4. **External Stakeholders**: Consider partners, vendors, regulators, and other external parties
            5. **Requirements Analysis**: Research functional and non-functional requirements from stakeholder perspectives
            6. **Change Management**: Analyze adoption challenges and change management strategies
            
            Project context: {project_context}
            Organization type: {organization}
            """
        }

    def _get_project_client(self):
        project_endpoint = os.environ["PROJECT_ENDPOINT"]
        return AIProjectClient(
            endpoint=project_endpoint,
            credential=DefaultAzureCredential(),
        )

    def _fetch_and_print_new_agent_response(self, thread_id: str, agents_client: AgentsClient, last_message_id: Optional[str] = None) -> Optional[str]:
        response = agents_client.messages.get_last_message_by_role(
            thread_id=thread_id,
            role=MessageRole.AGENT,
        )
        if not response or response.id == last_message_id:
            return last_message_id

        return response.id

    def deep_research(self, query: str) -> str:
        """Perform deep research on a topic using Azure AI agents with Bing grounding.
        
        Args:
            query: The research question or topic to investigate.
            
        Returns:
            str: Comprehensive research summary with citations.
        """
        try:
            with self._get_project_client() as project_client:
                conn_id = project_client.connections.get(name=os.environ["BING_RESOURCE_NAME"]).id
                
                deep_research_tool = DeepResearchTool(
                    bing_grounding_connection_id=conn_id,
                    deep_research_model=os.environ["DEEP_RESEARCH_MODEL_DEPLOYMENT_NAME"],
                )

                with project_client.agents as dr_agents_client:
                    dr_agent = dr_agents_client.create_agent(
                        model=os.environ["MODEL_DEPLOYMENT_NAME"],
                        name="research-agent",
                        instructions="You are a helpful Agent that assists in researching topics comprehensively.",
                        tools=deep_research_tool.definitions,
                    )

                    thread = dr_agents_client.threads.create()
                    
                    message = dr_agents_client.messages.create(
                        thread_id=thread.id,
                        role="user",
                        content=query,
                    )

                    run = dr_agents_client.runs.create(thread_id=thread.id, agent_id=dr_agent.id)
                    last_message_id = None
                    
                    while run.status in ("queued", "in_progress"):
                        time.sleep(2)
                        run = dr_agents_client.runs.get(thread_id=thread.id, run_id=run.id)
                        last_message_id = self._fetch_and_print_new_agent_response(
                            thread_id=thread.id,
                            agents_client=dr_agents_client,
                            last_message_id=last_message_id,
                        )

                    if run.status == "failed":
                        return f"Research failed: {run.last_error}"

                    final_message = dr_agents_client.messages.get_last_message_by_role(
                        thread_id=thread.id, role=MessageRole.AGENT
                    )
                    
                    if final_message:
                        # Format the response
                        text_summary = "\n\n".join([t.text.value.strip() for t in final_message.text_messages])
                        
                        # Add citations if present
                        if final_message.url_citation_annotations:
                            text_summary += "\n\n## References\n"
                            seen_urls = set()
                            for ann in final_message.url_citation_annotations:
                                url = ann.url_citation.url
                                title = ann.url_citation.title or url
                                if url not in seen_urls:
                                    text_summary += f"- [{title}]({url})\n"
                                    seen_urls.add(url)
                        
                        # Clean up
                        dr_agents_client.delete_agent(dr_agent.id)
                        return text_summary
                    else:
                        dr_agents_client.delete_agent(dr_agent.id)
                        return "No research results returned"

        except Exception as e:
            return f"Deep research failed: {e}"
    
    def structured_research(self, research_type: str, project_context: dict) -> str:
        """Perform structured research using predefined templates.
        
        Args:
            research_type: Type of research (technical, market, risk, stakeholder)
            project_context: Dictionary with project details
            
        Returns:
            str: Research results formatted as markdown
        """
        if research_type not in self.research_templates:
            return f"Unknown research type: {research_type}. Available types: {list(self.research_templates.keys())}"
        
        # Format the template with project context
        template = self.research_templates[research_type]
        try:
            formatted_query = template.format(**project_context)
            return self.deep_research(formatted_query)
        except KeyError as e:
            return f"Missing required project context field: {e}"
        except Exception as e:
            return f"Structured research failed: {e}"


class ProjectKickoffTool:
    def __init__(self):
        self.ado_tool = AzureDevOpsMCPTool()
        self.research_tool = DeepResearchMCPTool()
    
    def research_project_kickoff(self, story_id: str, research_types: List[str] = None) -> Dict[str, str]:
        """Perform comprehensive project kickoff research.
        
        Args:
            story_id: Azure DevOps story ID
            research_types: List of research types to perform
            
        Returns:
            Dict with research results for each type
        """
        if research_types is None:
            research_types = ["technical", "market", "risk", "stakeholder"]
        
        try:
            # Get project details from ADO
            story_details = self.ado_tool.get_azure_devops_story(story_id)
            
            # Extract project context from story
            project_context = self._extract_project_context(story_details, story_id)
            
            # Perform research for each type (limit to 2 types to avoid timeout)
            if len(research_types) > 2:
                research_types = research_types[:2]  # Limit to first 2 types for performance
            
            research_results = {}
            for research_type in research_types:
                if research_type in self.research_tool.research_templates:
                    print(f"Starting {research_type} research...", file=sys.stderr)
                    research_results[research_type] = self.research_tool.structured_research(
                        research_type, project_context
                    )
                    print(f"Completed {research_type} research", file=sys.stderr)
                else:
                    research_results[research_type] = f"Unknown research type: {research_type}"
            
            return {
                "project_details": story_details,
                "research_results": research_results,
                "project_context": project_context
            }
            
        except Exception as e:
            return {"error": f"Project kickoff research failed: {e}"}
    
    def _extract_project_context(self, story_details: str, story_id: str) -> dict:
        """Extract project context from ADO story details."""
        # Parse story details to extract key information
        lines = story_details.split('\n')
        
        # Extract title (assuming format: **Story ID: Title**)
        title = "Unknown Project"
        for line in lines:
            if line.startswith("**Story") and ":" in line:
                title = line.split(":", 1)[1].strip().rstrip("*")
                break
        
        # Extract description
        description = ""
        in_description = False
        for line in lines:
            if "**Description:**" in line:
                in_description = True
                continue
            if in_description:
                description += line + " "
        
        return {
            "project_name": title,
            "project_context": description.strip() or f"Azure DevOps Story {story_id}",
            "story_id": story_id,
            "technology_stack": "To be determined",  # Could be extracted from description
            "domain": "Enterprise Software",  # Default, could be extracted
            "industry": "Technology",  # Default, could be extracted
            "timeline": "To be determined",  # Could be extracted from dates
            "organization": "Enterprise"  # Default
        }


class DocumentGenerationTool:
    def __init__(self):
        pass
    
    def _get_blob_service_client(self):
        """Get Azure Blob Storage client."""
        connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
        
        if connection_string:
            # Check if it's a SAS URL or connection string
            if connection_string.startswith("https://") and "?" in connection_string:
                # It's a SAS URL, extract account URL and use SAS token
                account_url = connection_string.split("?")[0].rsplit("/", 1)[0]  # Remove container and SAS params
                return BlobServiceClient(account_url=account_url)
            else:
                # It's a proper connection string
                return BlobServiceClient.from_connection_string(connection_string)
        else:
            # Use DefaultAzureCredential if connection string not available
            account_url = f"https://{os.getenv('AZURE_STORAGE_ACCOUNT_NAME', 'defaultaccount')}.blob.core.windows.net"
            return BlobServiceClient(account_url=account_url, credential=DefaultAzureCredential())
    
    def generate_document(self, content: dict, document_type: str = "word", 
                         storage_container: str = "projects") -> str:
        """Generate and store a project kickoff document.
        
        Args:
            content: Dictionary with research results and project details
            document_type: Type of document to generate (word, markdown)
            storage_container: Azure storage container name
            
        Returns:
            str: URL or path to the generated document
        """
        try:
            if document_type == "word":
                return self._generate_word_document(content, storage_container)
            elif document_type == "markdown":
                return self._generate_markdown_document(content, storage_container)
            else:
                return f"Unsupported document type: {document_type}"
                
        except Exception as e:
            return f"Document generation failed: {e}"
    
    def _generate_word_document(self, content: dict, container: str) -> str:
        """Generate Word document from research content."""
        doc = Document()
        
        # Add title
        project_name = content.get("project_context", {}).get("project_name", "Project Kickoff Research")
        title = doc.add_heading(f"{project_name} - Kickoff Research", 0)
        
        # Add project details section
        doc.add_heading("Project Overview", level=1)
        if "project_details" in content:
            doc.add_paragraph(content["project_details"])
        
        # Add research results
        if "research_results" in content:
            for research_type, results in content["research_results"].items():
                doc.add_heading(f"{research_type.title()} Research", level=1)
                doc.add_paragraph(results)
                doc.add_page_break()
        
        # Add generation timestamp
        doc.add_paragraph(f"\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Save to blob storage
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"project-kickoff-{content.get('project_context', {}).get('story_id', 'unknown')}-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
        
        return self._upload_to_blob_storage(buffer.getvalue(), filename, container, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    def _generate_markdown_document(self, content: dict, container: str) -> str:
        """Generate Markdown document from research content."""
        project_name = content.get("project_context", {}).get("project_name", "Project Kickoff Research")
        
        md_content = f"# {project_name} - Kickoff Research\n\n"
        
        # Add project details
        md_content += "## Project Overview\n\n"
        if "project_details" in content:
            md_content += content["project_details"] + "\n\n"
        
        # Add research results
        if "research_results" in content:
            for research_type, results in content["research_results"].items():
                md_content += f"## {research_type.title()} Research\n\n"
                md_content += results + "\n\n"
        
        # Add generation timestamp
        md_content += f"\n*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n"
        
        filename = f"project-kickoff-{content.get('project_context', {}).get('story_id', 'unknown')}-{datetime.now().strftime('%Y%m%d-%H%M%S')}.md"
        
        return self._upload_to_blob_storage(md_content.encode('utf-8'), filename, container, "text/markdown")
    
    def _upload_to_blob_storage(self, content: bytes, filename: str, container: str, content_type: str) -> str:
        """Upload content to Azure Blob Storage."""
        try:
            connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
            
            if connection_string and connection_string.startswith("https://") and "?" in connection_string:
                # Handle SAS URL directly
                from azure.storage.blob import BlobClient
                sas_url = connection_string
                # Replace container in SAS URL if different
                if f"/{container}?" not in sas_url:
                    base_url = sas_url.split("?")[0].rsplit("/", 1)[0]  # Remove container from URL
                    sas_params = sas_url.split("?", 1)[1]  # Get SAS parameters
                    sas_url = f"{base_url}/{container}?{sas_params}"
                
                blob_client = BlobClient.from_blob_url(f"{sas_url.split('?')[0]}/{filename}?{sas_url.split('?')[1]}")
            else:
                # Use blob service client
                blob_service_client = self._get_blob_service_client()
                blob_client = blob_service_client.get_blob_client(container=container, blob=filename)
            
            from azure.storage.blob import ContentSettings
            content_settings = ContentSettings(content_type=content_type)
            blob_client.upload_blob(content, overwrite=True, content_settings=content_settings)
            
            return f"Document uploaded successfully: {blob_client.url.split('?')[0]}"  # Remove SAS params from returned URL
            
        except Exception as e:
            return f"Failed to upload to blob storage: {e}"